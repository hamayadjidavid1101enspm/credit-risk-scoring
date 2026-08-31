from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

# Initialisation du document
doc = Document()

# Titre principal
title = doc.add_heading('Rapport Complet du Projet PAEI : Modélisation du Risque de Crédit', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Auteur : David Janvion HAMAYADJI NGOMNA').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Encadrant : Ingénieur Archange KOUMBA MOUITY\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Architecture et Données
add_heading(doc, '1. Présentation des Données et Architecture du Projet', 1)
doc.add_paragraph(
    "L'objectif de ce projet est de prédire la probabilité de défaut de paiement d'un client. "
    "Le jeu de données utilisé provient de Kaggle ('Loan_default.csv') et contient plus de 250 000 profils "
    "de clients avec 18 variables (Âge, Revenu, Score de Crédit, etc.).\n"
    "L'architecture du projet a été pensée pour être professionnelle :\n"
    "- data/ : Données brutes et traitées.\n"
    "- src/ : Scripts Python automatisés (Pipeline).\n"
    "- models/ : Sauvegarde des algorithmes (fichiers .pkl et .json).\n"
    "- app/ : Interface Web pour la prise de décision."
)

# 2. Analyse Exploratoire
add_heading(doc, '2. Analyse Exploratoire des Données (EDA)', 1)
doc.add_paragraph("Avant de modéliser, une analyse statistique a été menée pour comprendre le comportement des variables.")

doc.add_heading('2.1 Analyse Univariée', 2)
doc.add_paragraph("L'analyse univariée permet de vérifier la distribution de chaque variable dans notre population. Par exemple, les revenus montrent une asymétrie à droite classique en finance (beaucoup de salaires moyens, peu de très hauts salaires).")
try:
    doc.add_picture('../reports/figures/univar_numeriques.png', width=Inches(5.5))
except:
    doc.add_paragraph("[Image univar_numeriques.png introuvable]")

doc.add_heading('2.2 Analyse Bivariée', 2)
doc.add_paragraph("Les boxplots croisent nos variables avec la cible (Défaut). On observe visuellement que les clients en défaut (1) ont une médiane de Taux d'Intérêt beaucoup plus haute et une médiane de Revenu plus basse.")
try:
    doc.add_picture('../reports/figures/bivar_numeriques.png', width=Inches(5.5))
except:
    pass

doc.add_heading('2.3 Matrice de Corrélation', 2)
doc.add_paragraph("La matrice de Pearson mesure l'intensité mathématique des liaisons. Le taux d'intérêt a la corrélation positive la plus forte avec le défaut. Il n'y a pas de multicolinéarité majeure entre les variables.")
try:
    doc.add_picture('../reports/figures/correlation_matrix.png', width=Inches(5.5))
except:
    pass

# 3. Justification des choix algorithmiques
add_heading(doc, '3. Justification des Choix Algorithmiques', 1)
doc.add_paragraph(
    "Au vu de la complexité des données financières, deux familles d'algorithmes ont été sélectionnées pour des raisons très précises :"
)
doc.add_heading('A. Pourquoi le Random Forest (Forêt Aléatoire) ?', 2)
doc.add_paragraph(
    "Nous avons utilisé un Random Forest dans la phase de recherche pour deux raisons majeures :\n"
    "1. Robustesse aux valeurs extrêmes : Les modèles basés sur des arbres de décision gèrent très bien les données non-linéaires et ne nécessitent pas une normalisation parfaite.\n"
    "2. Feature Importance : Ce modèle est excellent pour extraire 'l'importance des variables' (quels critères fendent l'arbre de décision). Il a permis de prouver que le Taux d'Intérêt et le Revenu étaient les 2 piliers du risque."
)
try:
    doc.add_picture('../reports/figures/feature_importance_rf.png', width=Inches(5.5))
except:
    pass

doc.add_heading('B. Pourquoi la Régression Logistique pour la Production ?', 2)
doc.add_paragraph(
    "Bien que le Random Forest soit puissant, nous avons choisi la Régression Logistique pour le simulateur final (l'application web) car le milieu bancaire exige une Transparence et Interprétabilité totale. "
    "La régression logistique fournit des poids mathématiques exacts (coefficients) qui peuvent être extraits (au format JSON) et injectés directement dans le code source d'un site Web (Javascript) pour un calcul instantané, sans nécessiter de serveur lourd."
)

# 4. Construction de l'Indice de Risque
add_heading(doc, '4. Méthodologie : Indice de Risque', 1)
doc.add_paragraph("Nous avons comparé deux approches pour construire le score de risque.")
doc.add_heading('Approche 1 : ACP (Analyse en Composantes Principales)', 2)
doc.add_paragraph("L'ACP réduit les dimensions pour créer un indice géométrique. Cependant, la séparation entre les bons et mauvais payeurs reste imparfaite à cause de la linéarité stricte de l'ACP.")
try:
    doc.add_picture('../reports/figures/risk_index_pca.png', width=Inches(5.5))
except:
    pass

doc.add_heading('Approche 2 : Machine Learning', 2)
doc.add_paragraph("En extrayant la probabilité (predict_proba) du Random Forest, nous obtenons un indice qui sépare massivement les clients risqués vers la droite. Cela permet à la banque de fixer un seuil (Cut-off) très clair.")
try:
    doc.add_picture('../reports/figures/risk_index_ml.png', width=Inches(5.5))
except:
    pass

# 5. Évaluation et AUC-ROC
add_heading(doc, "5. Évaluation du Modèle : Comprendre l'AUC et la courbe ROC", 1)
doc.add_paragraph(
    "Pour évaluer rigoureusement le modèle, les données ont été divisées : 80% pour l'entraînement (Train) "
    "et 20% pour le test (Test). Le modèle a été évalué sur des données qu'il n'avait jamais vues.\n\n"
    "Les résultats ont été mesurés avec la métrique AUC (Area Under the ROC Curve). La courbe ROC trace le "
    "taux de Vrais Positifs (bons signalements de défaut) contre le taux de Faux Positifs (fausses alertes). "
    "Notre modèle a obtenu un score AUC d'environ 0.76. Concrètement, cela signifie que si l'on prend un mauvais "
    "payeur au hasard et un bon payeur au hasard, notre Intelligence Artificielle classera le mauvais payeur comme "
    "étant plus risqué dans 76% des cas (contre 50% pour le hasard pur). C'est un score très robuste pour anticiper le comportement humain."
)

# 6. Mise en production
add_heading(doc, '6. Mise en Production (Simulateur Web)', 1)
doc.add_paragraph(
    "Le pipeline complet s'achève par l'automatisation logicielle. Le fichier 'main_pipeline.py' gère le feature "
    "engineering (création du ratio Prêt/Revenu), la normalisation (StandardScaler) et l'exportation des poids JSON. "
    "Ces poids sont injectés dans le fichier 'simulateur_scoring_credit.html', créant un Dashboard interactif et professionnel "
    "pour le conseiller bancaire, avec jauge de risque dynamique et décision automatisée."
)

# Sauvegarde
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Rapport_Complet_PAEI.docx')
doc.save(output_path)
print(f"Document Word sauvegardé avec succès à : {output_path}")
