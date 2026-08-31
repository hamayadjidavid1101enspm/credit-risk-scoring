import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Titre principal
title = doc.add_heading('Rapport Complet du Projet PAEI : Modélisation du Risque de Crédit', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Auteur : David Janvion HAMAYADJI NGOMNA').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Encadrant : Ingénieur Archange KOUMBA MOUITY\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Dossier absolu des images
IMG_DIR = r"C:\Users\DAVID PC\Documents\credit-risk-scoring\reports\figures"

# Fonction utilitaire pour insérer texte et image
def add_section_with_image(doc, title, text, image_name, interpretation):
    doc.add_heading(title, 2)
    doc.add_paragraph(text)
    
    img_path = os.path.join(IMG_DIR, image_name)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
    else:
        doc.add_paragraph(f"[Erreur: Image introuvable à {img_path}]")
        
    p = doc.add_paragraph()
    r = p.add_run("Interprétation : ")
    r.bold = True
    p.add_run(interpretation)
    doc.add_paragraph("-" * 20)

# 1. Architecture et Données
doc.add_heading('1. Présentation des Données et Architecture du Projet', 1)
doc.add_paragraph(
    "L'objectif de ce projet est de prédire la probabilité de défaut de paiement d'un client. "
    "Le jeu de données utilisé provient de Kaggle ('Loan_default.csv') et contient plus de 250 000 profils "
    "de clients avec 18 variables (Âge, Revenu, Score de Crédit, etc.).\n"
    "L'architecture du projet a été pensée pour être professionnelle :\n"
    "- data/ : Données brutes et traitées.\n"
    "- src/ : Scripts Python automatisés (Pipeline).\n"
    "- models/ : Sauvegarde des algorithmes.\n"
    "- app/ : Interface Web pour la prise de décision."
)

# 2. Analyse Exploratoire
doc.add_heading('2. Analyse Exploratoire des Données (EDA)', 1)
doc.add_paragraph("Avant de modéliser, une analyse statistique a été menée pour comprendre le comportement des variables.")

add_section_with_image(
    doc,
    '2.1 Analyse Univariée',
    "L'analyse univariée permet de vérifier la distribution individuelle de chaque variable financière et démographique.",
    'univar_numeriques.png',
    "Les histogrammes montrent la courbe de répartition globale de nos clients. Par exemple, la courbe des revenus ('Income') présente une asymétrie vers la droite (une majorité de salaires moyens et peu de hauts salaires). Cela valide la représentativité et la qualité de notre jeu de données, qui reflète la réalité économique."
)

add_section_with_image(
    doc,
    '2.2 Analyse Bivariée',
    "Les boxplots confrontent chaque variable avec la variable cible (Default : 0 = Sain, 1 = Défaut).",
    'bivar_numeriques.png',
    "C'est la preuve visuelle du risque. On observe clairement que la médiane (ligne du milieu de la boîte) du Taux d'Intérêt ('InterestRate') est nettement plus haute chez les clients en défaut (boîte orange). Inversement, la médiane des revenus est plus basse. Ces variables séparent donc naturellement les profils à risque."
)

add_section_with_image(
    doc,
    '2.3 Matrice de Corrélation',
    "Calcul des coefficients de Pearson pour mesurer l'intensité mathématique des liaisons.",
    'correlation_matrix.png',
    "La matrice confirme les graphiques précédents. Sur la ligne de la cible 'Default', le taux d'intérêt a la corrélation positive la plus forte. De plus, on n'observe aucune forte corrélation entre les variables d'entrée elles-mêmes, ce qui prouve l'absence de colinéarité (le modèle ne sera pas faussé)."
)

# 3. Justification algorithmique
doc.add_heading('3. Justification des Choix Algorithmiques', 1)

add_section_with_image(
    doc,
    'A. Pourquoi le Random Forest (Forêt Aléatoire) ?',
    "Nous avons utilisé un Random Forest dans la phase de recherche car les modèles basés sur des arbres gèrent parfaitement les données non-linéaires et sont robustes aux valeurs extrêmes. Surtout, ils permettent d'extraire la 'Feature Importance' de manière très précise.",
    'feature_importance_rf.png',
    "L'algorithme confirme notre analyse humaine. Les 3 variables clés qui déclenchent mathématiquement un défaut sont : 1. Le Taux d'Intérêt, 2. Le Revenu, 3. Le Montant du prêt. Les données démographiques (Mariage, Éducation) ont un impact très mineur. Nous nous sommes appuyés sur ce constat pour bâtir l'Indice de Risque."
)

doc.add_heading('B. Pourquoi la Régression Logistique pour la Production ?', 2)
doc.add_paragraph(
    "Pour le simulateur final (l'application web), nous avons opté pour la Régression Logistique. Le milieu bancaire exige une Transparence totale. "
    "La régression logistique fournit des coefficients mathématiques simples qui peuvent être exportés (via un fichier JSON) et intégrés au code source d'un site Web. Cela permet d'obtenir une probabilité nette sans dépendre d'un serveur lourd."
)

# 4. Indices
doc.add_heading('4. Méthodologie : Indice de Risque', 1)

add_section_with_image(
    doc,
    'Approche 1 : ACP (Analyse en Composantes Principales)',
    "L'ACP condense la variance financière sur un seul axe pour créer un premier indice (0 à 100).",
    'risk_index_pca.png',
    "On observe que la séparation n'est pas parfaite (les courbes se chevauchent au centre). L'ACP étant une approche strictement linéaire, elle peine à capturer la complexité du comportement humain pour séparer efficacement les profils."
)

add_section_with_image(
    doc,
    'Approche 2 : Machine Learning',
    "Extraction de la probabilité générée par le Random Forest (predict_proba).",
    'risk_index_ml.png',
    "Ici, la séparation est excellente ! La courbe des clients en défaut (orange) est fortement poussée vers les scores élevés (à droite). Grâce à cette distribution, la banque peut fixer une règle de coupure (Cut-off) claire pour refuser un prêt si le score ML dépasse un certain seuil."
)

# 5. AUC-ROC
doc.add_heading("5. Évaluation du Modèle : Comprendre l'AUC et la courbe ROC", 1)
doc.add_paragraph(
    "Pour évaluer le modèle sans tricherie, les données ont été divisées : 80% pour l'entraînement "
    "et 20% pour le test (données totalement inconnues du modèle).\n\n"
    "Les résultats ont été mesurés avec la métrique AUC-ROC (Area Under the Receiver Operating Characteristic Curve). "
    "Cette courbe trace le compromis entre les bons signalements (Vrais Positifs) et les fausses alertes (Faux Positifs). "
    "Notre modèle a obtenu un score AUC d'environ 0.76. Métierement parlant, cela signifie que si l'on prend un mauvais "
    "payeur et un bon payeur au hasard, l'IA a 76% de chances de correctement classer le mauvais payeur comme "
    "étant le plus risqué. C'est un score très performant pour le scoring de crédit."
)

# 6. Production
doc.add_heading('6. Mise en Production (Simulateur Web)', 1)
doc.add_paragraph(
    "Le pipeline complet s'achève par l'automatisation logicielle ('main_pipeline.py'). "
    "Les poids générés sont injectés dans 'simulateur_scoring_credit.html', créant un Dashboard interactif "
    "avec une jauge de risque dynamique pour un calcul instantané."
)

# Sauvegarde
output_path = r"C:\Users\DAVID PC\Documents\credit-risk-scoring\src\Rapport_Complet_PAEI_v2.docx"
doc.save(output_path)
print(f"Fichier écrasé et sauvegardé à {output_path}")
